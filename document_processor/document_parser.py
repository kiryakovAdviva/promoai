# --- START OF FILE document_parser.py ---
import os
import sys
import traceback
from typing import Any, Dict, List, Optional, Tuple, Generator

import pandas as pd
import pdfplumber
import docx
from tqdm import tqdm
# import openpyxl # Не нужен явный импорт, pandas его использует

# Используем относительные импорты
# Если запускаете не как пакет, замените на 'from common_utils import ...'
try:
    from .common_utils import clean_text, format_table_to_markdown
except ImportError:
    # Прямой импорт для случая плоской структуры файлов
    try:
        from common_utils import clean_text, format_table_to_markdown
        print("✅ common_utils imported directly into document_parser.")
    except ImportError as e:
        sys.stderr.write(f"❌ CRITICAL: Failed to import common_utils in document_parser. Error: {e}\n")
        sys.exit(1)

# --- Структура для возвращаемых "сырых" блоков ---
class RawContentBlock:
    """Структура для представления сырого блока контента из документа."""
    def __init__(self, type: str, content: Any, source_info: Dict):
        self.type: str = type # 'text', 'table', 'excel_row'
        self.content: Any = content # str для текста, List[Dict] для таблицы, Dict для строки Excel
        self.source_info: Dict = source_info # page_number, element_index, headers, etc.

    def __repr__(self):
        content_len_str = 'N/A'
        if isinstance(self.content, (str, list, dict)):
            try:
                content_len_str = str(len(self.content))
            except TypeError:
                pass # Не у всех объектов есть len
        return (f"RawContentBlock(type='{self.type}', "
                f"source='{self.source_info.get('document_name', 'Unknown')}', "
                f"page={self.source_info.get('page_number', 'N/A')}, "
                f"content_len={content_len_str})")

# --- PDF Parser ---
def parse_pdf(file_path: str) -> Generator[RawContentBlock, None, None]:
    """Парсит PDF файл."""
    try:
        with pdfplumber.open(file_path) as pdf:
            # Собираем все ссылки из документа
            all_hyperlinks = []
            for page in pdf.pages:
                try:
                    links = page.hyperlinks
                    if links:
                        all_hyperlinks.extend(links)
                except Exception as e:
                    print(f"  ⚠️ Error extracting links from page: {str(e)}")
                    continue
            
            print(f"  🔗 PDF Links found in '{os.path.basename(file_path)}': {len(all_hyperlinks)}")
            
            # Обрабатываем каждую страницу
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # Извлекаем текст
                    text = page.extract_text()
                    if text:
                        yield RawContentBlock(
                            content=text,
                            type='text',
                            source_info={
                                'page_number': page_num,
                                'document_hyperlinks': all_hyperlinks,
                                'document_name': os.path.basename(file_path)
                            }
                        )
                    
                    # Извлекаем таблицы
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            if not table or not any(any(cell for cell in row) for row in table):
                                print(f"    ℹ️ PDF Table on p.{page_num} has headers but no data. Skipping.")
                                continue
                                
                            # Очищаем таблицу от None значений
                            cleaned_table = []
                            for row in table:
                                cleaned_row = [str(cell).strip() if cell is not None else '' for cell in row]
                                if any(cleaned_row):  # Пропускаем полностью пустые строки
                                    cleaned_table.append(cleaned_row)
                            
                            if not cleaned_table:
                                continue
                                
                            # Определяем заголовки таблицы
                            headers = cleaned_table[0] if cleaned_table else []
                            
                            # Если первая строка не похожа на заголовки, генерируем их
                            if not any(h.strip() for h in headers):
                                headers = [f"Column_{i+1}" for i in range(len(cleaned_table[0]))]
                                table_data = cleaned_table
                            else:
                                table_data = cleaned_table[1:]
                            
                            yield RawContentBlock(
                                content=table_data,
                                type='table',
                                source_info={
                                    'page_number': page_num,
                                    'table_number': table_num,
                                    'headers': headers,
                                    'document_hyperlinks': all_hyperlinks,
                                    'document_name': os.path.basename(file_path)
                                }
                            )
                            
                except Exception as e:
                    print(f"  ⚠️ Error processing page {page_num}: {str(e)}")
                    continue
                    
    except Exception as e:
        print(f"❌ Error opening PDF file: {str(e)}")
        raise

# --- DOCX Parser ---
def parse_docx(file_path: str) -> Generator[RawContentBlock, None, None]:
    """Парсит DOCX файл."""
    try:
        doc = docx.Document(file_path)
        
        # Собираем все ссылки из документа
        all_hyperlinks = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//w:hyperlink'):
                    all_hyperlinks.append({
                        'text': run.text,
                        'url': run._element.xpath('.//w:hyperlink/@r:id')[0]
                    })
        
        # Обрабатываем параграфы
        current_heading = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Определяем тип параграфа
            if paragraph.style.name.startswith('Heading'):
                current_heading = text
                yield RawContentBlock(
                    content=text,
                    type='heading',
                    source_info={
                        'heading_level': int(paragraph.style.name[-1]),
                        'document_hyperlinks': all_hyperlinks
                    }
                )
            else:
                yield RawContentBlock(
                    content=text,
                    type='text',
                    source_info={
                        'current_heading': current_heading,
                        'document_hyperlinks': all_hyperlinks
                    }
                )
        
        # Обрабатываем таблицы
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            headers = []
            
            for row_num, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if row_num == 0:
                    headers = row_data
                table_data.append(row_data)
            
            if table_data:
                yield RawContentBlock(
                    content=table_data,
                    type='table',
                    source_info={
                        'table_number': table_num,
                        'headers': headers,
                        'document_hyperlinks': all_hyperlinks
                    }
                )
                
    except Exception as e:
        print(f"❌ Error opening DOCX file: {str(e)}")
        raise

# --- Excel Parser (v17 - Refined, using RawContentBlock) ---
def parse_excel(xlsx_path: str) -> Generator[RawContentBlock, None, None]:
    """Парсит Excel файл."""
    try:
        print(f"\n📊 Processing Excel file: {os.path.basename(xlsx_path)}")
        # Читаем все листы
        excel_file = pd.ExcelFile(xlsx_path)
        
        for sheet_name in excel_file.sheet_names:
            print(f"  📑 Processing sheet: {sheet_name}")
            # Читаем с поддержкой пустых ячеек
            df = pd.read_excel(xlsx_path, sheet_name=sheet_name, dtype=str)
            
            # Заменяем NaN на пустые строки
            df = df.fillna('')
            
            # Проверяем, есть ли данные в таблице
            if df.empty:
                print(f"    ⚠️ Sheet '{sheet_name}' is empty")
                continue
                
            # Проверяем, есть ли непустые строки
            non_empty_rows = df.apply(lambda row: any(str(cell).strip() for cell in row), axis=1)
            if not any(non_empty_rows):
                print(f"    ⚠️ Sheet '{sheet_name}' contains only empty rows")
                continue
            
            print(f"    ✅ Found {len(df)} rows in sheet '{sheet_name}'")
            
            # Получаем заголовки
            headers = df.columns.tolist()
            
            # Создаем один чанк для всего листа
            table_data = []
            for idx, row in df.iterrows():
                # Пропускаем пустые строки
                if not any(str(cell).strip() for cell in row):
                    print(f"    ℹ️ Row {idx + 1} is empty, skipping")
                    continue
                    
                # Преобразуем строку в список значений
                row_values = [str(cell).strip() for cell in row]
                table_data.append(row_values)
            
            if table_data:
                yield RawContentBlock(
                    type='table',
                    content=table_data,
                    source_info={
                        'headers': headers,
                        'sheet_name': sheet_name,
                        'document_name': os.path.basename(xlsx_path)
                    }
                )
                
    except Exception as e:
        print(f"❌ Error processing Excel file: {str(e)}")
        traceback.print_exc()

# --- Главная функция-диспетчер ---
def parse_document(file_path: str) -> Generator[RawContentBlock, None, None]:
    """
    Парсит документ и возвращает генератор блоков контента.
    Поддерживает PDF и DOCX файлы.
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        yield from parse_pdf(file_path)
    elif file_ext == '.docx':
        yield from parse_docx(file_path)
    elif file_ext == '.xlsx':
        yield from parse_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

# --- END OF FILE document_parser.py ---